"""Word (.docx) generation using python-docx."""
import os
from datetime import datetime, timezone
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_document(
    doc_id: str,
    plan: dict,
    sections: List[Dict[str, List[str]]],
    assumptions: List[str],
    reflection_notes: List[str],
    request: str,
    out_dir: str,
) -> str:
    """Render a structured business document and return its file path."""
    os.makedirs(out_dir, exist_ok=True)
    doc = Document()

    title = plan.get("title") or "Generated Document"
    doc_type = plan.get("document_type", "document")

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(doc_type.title())
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Metadata table
    meta = [
        ("Document type", doc_type.title()),
        ("Generated (UTC)", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")),
        ("Document ID", doc_id),
        ("Source request", request),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in meta:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)
        for para in cells[0].paragraphs:
            for r in para.runs:
                r.bold = True
    doc.add_paragraph()

    # Body sections
    for sec in sections:
        doc.add_heading(sec["name"], level=1)
        for para in sec.get("paragraphs", []):
            doc.add_paragraph(para)
        for bullet in sec.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    # Appendix: assumptions + agent self-check notes
    if assumptions:
        doc.add_heading("Assumptions", level=1)
        for item in assumptions:
            doc.add_paragraph(item, style="List Bullet")

    if reflection_notes:
        doc.add_heading("Agent Self-Check Notes", level=1)
        for note in reflection_notes:
            doc.add_paragraph(note, style="List Bullet")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run("Generated autonomously by RUKU")
    frun.italic = True
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    path = os.path.join(out_dir, f"{doc_id}.docx")
    doc.save(path)
    return path
